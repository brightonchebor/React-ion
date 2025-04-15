from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create a new Document
doc = Document()

# Add Header
header = doc.add_heading('BRIGHTON CHEBOR', 0)
header.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_heading('Backend Developer (Python)', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER

# Contact Info
doc.add_paragraph('Email: cheborbrighton805@gmail.com', style='Normal')
doc.add_paragraph('GitHub: github.com/brightonchebor', style='Normal')
doc.add_paragraph('LinkedIn: linkedin.com/in/brighton-chebor', style='Normal')
doc.add_paragraph('[Insert QR Code linking to GitHub]', style='Normal')

# Professional Summary
doc.add_heading('Professional Summary', level=2)
doc.add_paragraph(
    'A dedicated Backend Developer with expertise in Python, SQL, and REST API development. '
    'Proficient in building secure, scalable web applications using Django and Flask, with a '
    'strong focus on authentication, database management, and code quality. Passionate about '
    'leveraging technology to solve real-world problems and eager to contribute to innovative teams.'
)

# Skills
doc.add_heading('Skills', level=2)
skills = [
    'Languages: Python, SQL, JavaScript, HTML, CSS',
    'Frameworks: Django, Django REST Framework, Flask, Bootstrap',
    'Databases: MySQL, PostgreSQL, SQLite, Microsoft SQL Server',
    'Tools: Git, GitHub, Docker, Linux, Visual Studio Code, Postman, Mailtrap, Pytest',
    'Soft Skills: Problem-solving, Team Collaboration, Time Management'
]
for skill in skills:
    doc.add_paragraph(skill, style='ListBullet')

# Education
doc.add_heading('Education', level=2)
doc.add_paragraph('Pwani University', style='Normal')
doc.add_paragraph('BSc. Telecommunication and Information Technology', style='Normal')
doc.add_paragraph('September 2021 – Present', style='Normal')
doc.add_paragraph('Relevant Coursework: Data Structures, Web Development, Database Systems', style='Normal')

# Certification
doc.add_heading('Certification', level=2)
doc.add_paragraph('Certified Charity Donation Platform – eMobilis Kenya (2023)', style='Normal')
doc.add_paragraph(
    'Developed and deployed a web-based charity donation platform using Django, REST APIs, '
    'and payment integration. Implemented JWT-based authentication and integrated with local '
    'payment gateways. Bridged the gap between donors and small charity organizations in Kenya.',
    style='ListBullet'
)

# Page Break for Page 2
doc.add_page_break()

# Professional Experience
doc.add_heading('Professional Experience', level=2)
doc.add_paragraph('Freelance Backend Developer – Remote (June 2023 – Present)', style='Normal')
freelance = [
    'Designed and developed RESTful APIs for small businesses using Django REST Framework, improving client data management by 30%.',
    'Optimized database queries with PostgreSQL, reducing response times by 25%.',
    'Collaborated with frontend developers to integrate APIs with React-based interfaces.',
    'Tools: Django, PostgreSQL, Postman, Git'
]
for item in freelance:
    doc.add_paragraph(item, style='ListBullet')

doc.add_paragraph('Intern, eMobilis Kenya – Nairobi, Kenya (January 2023 – May 2023)', style='Normal')
intern = [
    'Contributed to the development of a charity donation platform, focusing on secure user authentication and API endpoints.',
    'Conducted unit testing with Pytest, achieving 90% code coverage.',
    'Managed database migrations and schema design using MySQL.',
    'Tools: Django, MySQL, Pytest, Docker'
]
for item in intern:
    doc.add_paragraph(item, style='ListBullet')

# Projects
doc.add_heading('Projects', level=2)
doc.add_paragraph('Charity Donation Platform (eMobilis Certification Project, 2023)', style='Normal')
charity = [
    'Built a centralized web platform to connect local donors with charities.',
    'Features: User authentication (JWT), payment integration, and admin dashboard.',
    'Tech Stack: Django, Django REST Framework, MySQL, Bootstrap',
    'Outcome: Deployed successfully, serving 10+ charities and 100+ donors.'
]
for item in charity:
    doc.add_paragraph(item, style='ListBullet')

doc.add_paragraph('Personal Blog API (Personal Project, 2024)', style='Normal')
blog = [
    'Developed a REST API for a blog platform with CRUD functionality.',
    'Implemented role-based authentication and content moderation.',
    'Tech Stack: Flask, SQLite, Postman',
    'Outcome: Hosted on GitHub, used as a learning tool for API development.'
]
for item in blog:
    doc.add_paragraph(item, style='ListBullet')

# Hobbies & Interests
doc.add_heading('Hobbies & Interests', level=2)
hobbies = [
    'Open-source contributions on GitHub',
    'Exploring cloud technologies (AWS, Azure)',
    'Blogging about Python and web development',
    'Playing chess and volunteering at local tech meetups'
]
for hobby in hobbies:
    doc.add_paragraph(hobby, style='ListBullet')

# Save the document
doc.save('Brighton_Chebor_CV.docx')